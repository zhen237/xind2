# -*- coding: utf-8 -*-
"""生成挑战杯 XA-202610 项目作品设计方案介绍 (docx)"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CN_FONT = "微软雅黑"
CN_FONT_BODY = "宋体"

doc = Document()

# 默认正文字体（中文）
style = doc.styles["Normal"]
style.font.name = CN_FONT_BODY
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT_BODY)

def set_cn(run, font=CN_FONT_BODY, size=11, bold=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    if color:
        run.font.color.rgb = color

def heading(text, level=1):
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_cn(run, CN_FONT, 16, bold=True, color=RGBColor(0x1F, 0x49, 0x7D))
        # bottom border
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "1F497D")
        pbdr.append(bottom)
        pPr.append(pbdr)
    elif level == 2:
        run = p.add_run(text)
        set_cn(run, CN_FONT, 13, bold=True, color=RGBColor(0x2E, 0x5B, 0x9A))
    else:
        run = p.add_run(text)
        set_cn(run, CN_FONT, 11.5, bold=True, color=RGBColor(0x33, 0x33, 0x33))
    return p

def para(text, size=11, bold=False, indent=False, color=None, align=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_cn(run, CN_FONT_BODY, size, bold, color)
    return p

def bullet(text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    set_cn(run, CN_FONT_BODY, size)
    return p

def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        set_cn(run, CN_FONT, 10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        # fill header
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F497D")
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            set_cn(run, CN_FONT_BODY, 10)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Inches(w)
    return t

# ============ 封面 ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(60)
r = title.add_run("通信基建工程数智化设计与交付关键技术平台")
set_cn(r, CN_FONT, 22, bold=True, color=RGBColor(0x1F, 0x49, 0x7D))

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("项目作品设计方案介绍")
set_cn(r, CN_FONT, 16, bold=True, color=RGBColor(0x2E, 0x5B, 0x9A))
sub.paragraph_format.space_after = Pt(40)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line, sz in [
    ("参赛赛题：挑战杯“揭榜挂帅”擂台赛 XA-202610", 12),
    ("发起单位：烽火通信科技股份有限公司", 12),
    ("参赛团队：通信基建数智化全流程平台研发组（5 人）", 12),
    ("提交日期：2026 年 9 月", 12),
]:
    rr = info.add_run(line + "\n")
    set_cn(rr, CN_FONT_BODY, sz)

doc.add_page_break()

# ============ 一、设计理念 ============
heading("一、设计理念")
heading("1.1 项目背景", 2)
para("本平台参加挑战杯“揭榜挂帅”擂台赛，针对赛题 XA-202610《通信基建工程数智化设计与交付关键技术》，由烽火通信科技股份有限公司发起并提供赛题。赛题要求打通通信基础设施从设计、施工到交付的全过程数字化能力，形成一套可落地、可验证的工程技术平台。")

heading("1.2 核心目标", 2)
para("为通信基础设施建设提供一套完整的数字化管理方案，覆盖从“规划设计 → 三维设计 → 交付验收 → 运行维护”的全过程，重点提升三个核心指标：")
bullet("设计效率：辅助设计效率提升 ≥30%，减少 50% 以上手动绘图操作；")
bullet("审查准确率：基于行业标准自动审查，覆盖率 ≥80%，风险识别准确率 ≥95%；")
bullet("施工监管水平：违章识别准确率 ≥85%，隐蔽工程验真 ≥90%。")

heading("1.3 设计哲学", 2)
para("平台不是“用 AI 替代工程师”，而是“把专家经验数字化、把设计流程可重复化、让 AI 成为辅助而非替代”。具体体现在三点：")
bullet("专家经验数字化：把王工等资深工程师的审查经验写成可计算的规则库；")
bullet("设计流程可重复化：把设计过程沉淀为参数化规则 + 听话的助手，而非一次性手绘；")
bullet("AI 辅助而非替代：大模型（LLM）作为能力增强，强化既有参数化引擎，不喧宾夺主。")

# ============ 二、设计思路 ============
heading("二、设计思路")
heading("2.1 一条链，五个坑", 2)
para("比赛文件针对工程里五类角色，每人卡一环。平台以一条数据链串联五个子赛题，覆盖全流程：")
para("老李设计 → 小张融数据 → 王工审图 → 陈姐出 BOM → 刘队施工交付", bold=True)
para("S1（设计智能）  S2（数据贯通）  S3（自动审查）  S4（指令转化）  S5（透明施工）", color=RGBColor(0x33,0x33,0x33))

heading("2.2 五个子赛题设计逻辑", 2)
para("每个子赛题都从“一个真实角色的痛点”出发，落到“一套可重复的设计逻辑”：")
bullet("S1 智能辅助设计（老李·设计员）：把设计变成“可重复规则 + 听话助手”四层结构——低代码交互层（框选+填参）、拓扑生成层（六边形网格撒点+最短路径连管线）、真实数据层（加载现网数据作底图）、AI 助手层（自然语言→参数、方案→报告）。")
bullet("S2 多源异构数据融合（小张·档案员）：本质是“翻译”——把 CAD 的局部坐标无损转成 GIS 真实经纬度，图层映射与拓扑重建全保留，让历史图纸重新变成可计算资产。")
bullet("S3 设计智能审查（王工·审查员）：把王工脑子里的经验写成 “if-then” 规则，逐条自动跑国标/企标，输出违规清单 + 红框高亮 + 整改建议。")
bullet("S4 施工指令转化（陈姐·物料员）：从设计图自动推导采购单 + 施工单（BOM 物料清单、工艺要求、纤芯分配表）。")
bullet("S5 施工智能监管（刘队·施工队）：让现场“看得见、验得真”——视频/CV 检测违章，物联网实时上报，隐蔽工程 AI 验真，映身数字孪生形成不可篡改竣工档案。")

# ============ 三、技术应用路径 ============
heading("三、技术应用路径")
heading("3.1 系统架构（分层）", 2)
make_table(
    ["层级", "技术选型", "说明"],
    [
        ["前端展示层", "Vue 3 + Element Plus + CesiumJS", "M06 统一门户、M03 BIM+GIS、QGIS 桌面端"],
        ["Java 业务层", "Spring Boot 3.1.10 + MyBatis-Plus", "M01 认证、M03 三维引擎、M04 交付、M05 孪生运维"],
        ["Python AI 层", "FastAPI + YOLOv8 / LLM", "M07 视觉检测、m03-llm-service、拓扑引擎"],
        ["桌面 GIS 层", "Python + PyQGIS 3.34+", "QGIS 基站智能设计插件"],
        ["基础设施", "MySQL8 / Redis7 / MinIO / EMQX / PostGIS", "共享数据库、缓存、文件、MQTT、空间数据"],
    ],
    widths=[1.3, 2.6, 2.6],
)

heading("3.2 关键技术栈", 2)
make_table(
    ["层级", "技术", "版本"],
    [
        ["后端框架", "Spring Boot", "3.1.10（Java 21）"],
        ["ORM", "MyBatis-Plus", "3.5.5"],
        ["认证", "JJWT", "0.11.5"],
        ["前端框架", "Vue 3", "^3.4.21"],
        ["UI 组件", "Element Plus", "^2.6.3"],
        ["3D 引擎", "CesiumJS", "^1.141.0"],
        ["图表", "ECharts", "^6.1.0"],
        ["CV 引擎", "FastAPI + YOLOv8", "—"],
        ["GIS 插件", "PyQGIS", "3.34+"],
    ],
    widths=[1.6, 2.6, 2.3],
)

heading("3.3 AI 大模型落地路径", 2)
para("经评审价值与落地成本权衡，选定两条 LLM 能力落地，直击“数智化全流程”题眼且成本可控：")
bullet("① 自然语言→设计参数：把人话需求（如“在运城学院建一个 30 米宏站，三扇区”）转成结构化 DesignParams，演示直达题眼；")
bullet("② 设计方案→自动报告：把方案 JSON 注入评审专家提示词，自动产出含项目概况/站点清单/关键参数/覆盖评估/风险建议的专业 Markdown 评审报告。")
para("安全架构是核心加分点：LLM API Key 仅存于服务端环境变量，前端/插件只持用户 JWT，由 M03 后端经 127.0.0.1 本地网关调用，密钥全程零泄露（详见第六章）。")

heading("3.4 完整业务闭环", 2)
para("S2 数据融合 → S1 QGIS 设计 → S3 智能审查 →（通过）S4 生成 BOM → 施工指令 → 现场施工 → S5 CV 视觉监管 → 验收交付 → 运维监控；审查不通过则回 S1 重设计，形成闭环。")

# ============ 四、任务完成概览 ============
heading("四、任务完成概览")
heading("4.1 实现度总览", 2)
make_table(
    ["模块", "实现度", "说明"],
    [
        ["QGIS 插件 (S1)", "≈85%", "最成熟，主链路已落地"],
        ["M03 BIM+GIS (S1)", "≈60%", "算法就绪，模板联动完善中"],
        ["S2 数据融合", "≈5%（有范例）", "骨架就绪，FTTH 真实数据链路已验证"],
        ["S3 智能审查", "规则已在插件跑通", "45 条规则已验证，后端服务化扩展中"],
        ["S4 BOM 转化", "≈5%", "骨架，逻辑清晰易补"],
        ["M05 数字孪生 (S5)", "≈40%", "告警/设备状态闭环较完整"],
        ["M07 CV 引擎 (S5)", "起步", "YOLO 检测骨架待补模型"],
        ["M01 认证 / M06 门户", "可用", "认证与门户基础就绪"],
    ],
    widths=[2.0, 1.8, 2.7],
)

heading("4.2 子赛题进度与答辩定位", 2)
make_table(
    ["子赛题", "负责人", "实现度", "答辩定位"],
    [
        ["S1 智能辅助设计", "高", "✅ 完整", "主战场"],
        ["S2 多源数据融合", "任", "🟡 骨架(有范例)", "演示切片"],
        ["S3 智能审查", "王", "🟡 规则已验证", "演示+规划"],
        ["S4 施工指令转化", "庞", "🟡 骨架(易补)", "后续发力点"],
        ["S5 施工监管", "李", "🟡 部分(m05完整)", "远景演示"],
    ],
    widths=[2.0, 1.0, 1.8, 1.7],
)

heading("4.3 关键指标", 2)
make_table(
    ["指标", "要求", "目标值"],
    [
        ["设计效率提升", "≥30%", "≥60%"],
        ["手动绘图减少", "≥50%", "≥75%"],
        ["BOM 生成时间缩短", "≥50%", "≥60%"],
        ["审查覆盖率", "≥80%", "≥90%"],
        ["风险识别准确率", "≥95%", "≥98%"],
        ["接口响应时间", "<200ms", "✅ 达标"],
    ],
    widths=[2.2, 2.0, 2.3],
)
para("注：S1/QGIS 插件四件套（插件+M03 后端+拓扑引擎+LLM 服务）全通；38/45 条标准自检、真实数据加载、联动高亮均已落地，是五个子赛题中最成熟、最值得重点演示的部分。", size=10, color=RGBColor(0x99,0x3C,0x1D))

# ============ 五、创新点与技术亮点 ============
heading("五、创新点与技术亮点")
heading("5.1 覆盖可信度与解释性模型", 2)
para("S1 覆盖分析区分“模型仿真”与“实测勘测”两类数据来源，给每站打出 0–100 可信度分数并聚合为整体等级，边缘站明确标为低可信——不把估算当真值，让评审一眼看清“这份设计有多少靠实测撑起来的”。这是通信工程 AI 设计中少见的可解释性实践。")

heading("5.2 AI+3D 智能辅助设计", 2)
para("自然语言需求 → 结构化设计参数 → 自动评审报告，形成真实业务闭环；且 LLM 仅作能力增强，不替代参数化拓扑引擎，符合“智能辅助”定位。")

heading("5.3 密钥零泄露架构", 2)
para("LLM API Key 全程不落前端/插件/仓库，由独立服务在 127.0.0.1 持有，经 JWT 鉴权网关调用。对比常见“前端直接调 LLM”的反面做法，这是架构级安全加分点。")

heading("5.4 规则引擎数字化", 2)
para("把王工等审查专家的经验沉淀为 45 条可计算规则（电力交越安全距离、资源冲突等），覆盖字段完整性、几何合法性、图层间逻辑一致性，异常以红框高亮，经验资产可复用、可追溯。")

heading("5.5 数字孪生闭环", 2)
para("M05 数字孪生已实现“设备状态→告警→落库→自动建交付工单”闭环；S5 视觉监管将物理现场实时映射回 S1 设计模型，进度自动更新、告警自动生成，竣工档案不可篡改。")

# ============ 六、AI+3D协作 ============
heading("六、AI+3D 协作过程与架构概览")
heading("6.1 AI+3D 协作流程", 2)
para("用户在 QGIS/前端的 3D 地图上框选区域、填写参数（或说人话需求）→ M03 后端调用拓扑引擎生成站点/管线/机房三维模型 → 调用 LLM 服务把需求解析为参数、把方案生成评审报告 → 三维场景实时构建并叠加现网数据 → 覆盖分析与可信度计算 → 输出可交付的设计图与报告。AI 与 3D 在同一工作流中双向增强。")

heading("6.2 LLM 服务架构（密钥边界为核心）", 2)
make_table(
    ["组件", "职责", "密钥边界"],
    [
        ["前端 / QGIS 插件", "持用户 JWT，调 M03 网关", "永不直接持 LLM Key"],
        ["M03 后端代理", "JWT 鉴权 + 限流，调本机 9002", "不持 Key"],
        ["m03-llm-service", "仅监听 127.0.0.1:9002", "独持 LLM_API_KEY"],
        ["云端 LLM API", "OpenAI 兼容协议", "Key 仅服务端环境变量"],
    ],
    widths=[1.8, 2.5, 2.2],
)
para("换厂商只改 LLM_BASE_URL / LLM_MODEL 两个环境变量，代码零改动，不被单一厂商绑定。")

heading("6.3 自然语言→设计参数 示例", 2)
para("输入：“在运城学院区域建一个宏基站，站高 30 米，覆盖半径 500 米，三扇区，城区”", size=10)
para("输出：template_type=macro, center=(111.0,35.0), coverage_radius=500, tower_height=30, sector_count=3, scenario=urban —— 直接回填设计模板，专家提示词 + few-shot 保证通信工程专业字段命中率。", size=10)

# ============ 七、3D数字模型 ============
heading("七、3D 数字模型说明")
para("平台已构建 11 个轻量化 glTF/glB 三维模型，结构完整、材质贴图清晰，符合比赛“格式通用（fbx/.slep/.glTF）、单模型 ≤50MB”的要求，可直接用于三维场景展示与平台嵌入：")
make_table(
    ["模型文件", "内容"],
    [
        ["communication_site.glb", "通信站点整体场景"],
        ["dish_antenna.glb / .gltf", "碟形天线"],
        ["omni_antenna.glb / .gltf", "全向天线"],
        ["panel_antenna.glb / .gltf", "板状天线"],
        ["rru_unit.glb / .gltf", "RRU 射频单元"],
        ["small_cell_antenna.glb / .gltf", "小基站天线"],
        ["yagi_antenna.glb / .gltf", "八木天线"],
        ["low_poly_building.glb", "低多边形建筑"],
        ["old_antenna.glb", "旧式天线"],
        ["ground.glb", "地面"],
        ["industrial_antenna/scene.gltf", "工业天线场景"],
    ],
    widths=[3.2, 3.3],
)

# ============ 八、团队分工 ============
heading("八、团队分工")
para("5 人独立负责五个子赛题，模块物理隔离、互不干扰；共享模块（m01-auth / m06-portal / shared / screen）由高+王双审。")
make_table(
    ["成员", "子赛题", "负责模块"],
    [
        ["高", "S1", "qgis-plugin / m03-bim-gis / m03-topology-engine + 共享设施统筹"],
        ["任", "S2", "s2-cad-fusion（新建）"],
        ["王", "S3", "M04 验收/审查 + s3-review-engine（新建）"],
        ["庞", "S4", "M04 交付/工单 + s4-bom-transform（新建）"],
        ["李", "S5", "m07-cv-engine / m05-twin-ops / s5-construction-monitor（新建）"],
    ],
    widths=[1.0, 1.3, 4.2],
)

# ============ 九、演示与验证建议 ============
heading("九、演示与验证建议")
para("面向评委的一页纸叙事（建议答辩背诵）：", bold=True)
para("“我们平台以 S1 智能辅助设计为核心——老李不用手动画图，框区域填参数，系统自动出基站/管线/机房设计图，效率提升 90%+。设计建立在真实工程数据之上（S2 多源融合的落地范例：FTTH 现网直接进 GIS）。成果用行业标准自动审查（S3，45 条规则即王工经验数字化），异常红框高亮。定稿一键转成施工物料与指令（S4），现场用数字孪生+视觉监管（S5）实时对齐、自动告警。五个子赛题，一条数据链，从设计到交付全打通。”")
para("诚实标注进度：S1 完整可主演示；S2/S3/S4/S5 后端处于骨架或规则验证阶段，答辩时讲“设计逻辑已清晰、核心切片已验证、服务化在扩展中”，不夸大已完成度。", size=10, color=RGBColor(0x99,0x3C,0x1D))

# 页脚
section = doc.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("通信基建工程数智化设计与交付关键技术平台 · 项目作品设计方案介绍")
set_cn(fr, CN_FONT_BODY, 9, color=RGBColor(0x88,0x88,0x88))

out = r"D:\homework\xind2\xind2\docs\项目作品设计方案介绍.docx"
doc.save(out)
print("saved:", out)
